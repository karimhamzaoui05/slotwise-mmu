from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "document_assets"
OUTPUT = DOCS / "SlotWise_MMU_Project_Documentation.docx"

TEAL = "087E8B"
DEEP_TEAL = "075F68"
DARK = "17252A"
MUTED = "5E737A"
CORAL = "FF6B5E"
LIGHT_TEAL = "E8F5F6"
LIGHT_GRAY = "F2F4F7"
BORDER = "DDE4E8"
WHITE = "FFFFFF"
GREEN = "16845B"
GOLD = "D98C10"
RED = "D64545"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=11, color=DARK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in edges.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_field(paragraph, code: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str, color=TEAL):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_numbering_definition(doc: Document, num_format: str, text: str, num_id: int, abstract_id: int):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_bullet(doc, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    apply_num(p, 31)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.16
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph()
    apply_num(p, 32)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.16
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_body(doc, text: str, bold_prefix: str | None = None, italic=False, color=DARK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic, color=color)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 11.5}[level],
        color={1: TEAL, 2: DEEP_TEAL, 3: DARK}[level],
        bold=True,
    )
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, left={"val": "single", "sz": "18", "color": accent})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_key_value_table(doc, rows: Sequence[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2250, 7110])
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        shade_cell(left, LIGHT_TEAL)
        shade_cell(right, WHITE)
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run_font(p1.add_run(label), size=10, color=DEEP_TEAL, bold=True)
        p2 = right.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        set_run_font(p2.add_run(value), size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_comparison_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, TEAL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(header), size=9.5, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows, start=1):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            shade_cell(cell, WHITE if row_idx % 2 else LIGHT_GRAY)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run_font(p.add_run(value), size=9.2, color=DARK)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def font_path(bold=False):
    paths = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for path in paths:
        if path.exists():
            return str(path)
    return None


def load_font(size: int, bold=False):
    path = font_path(bold)
    return ImageFont.truetype(path, size) if path else ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_arrow(draw, start, end, color=TEAL, width=5):
    draw.line([start, end], fill=f"#{color}", width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - 14 * direction, y2 - 9), (x2 - 14 * direction, y2 + 9)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - 14 * direction), (x2 + 9, y2 - 14 * direction)]
    draw.polygon(points, fill=f"#{color}")


def draw_box(draw, xy, title, subtitle=None, fill=WHITE, outline=TEAL, title_color=DARK):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=f"#{fill}", outline=f"#{outline}", width=3)
    title_font = load_font(24, True)
    small_font = load_font(18)
    lines = wrap_text(draw, title, title_font, x2 - x1 - 34)
    total = len(lines) * 29 + (24 if subtitle else 0)
    y = y1 + (y2 - y1 - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=title_font, fill=f"#{title_color}")
        y += 29
    if subtitle:
        y += 3
        bbox = draw.textbbox((0, 0), subtitle, font=small_font)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), subtitle, font=small_font, fill=f"#{MUTED}")


def create_architecture_diagram(path: Path):
    image = Image.new("RGB", (1500, 620), "white")
    draw = ImageDraw.Draw(image)
    title = load_font(30, True)
    draw.text((55, 35), "SlotWise MMU technical architecture", font=title, fill=f"#{DARK}")
    boxes = {
        "user": (55, 230, 280, 380),
        "ui": (380, 110, 675, 250),
        "state": (380, 365, 675, 505),
        "screens": (790, 90, 1120, 230),
        "models": (790, 275, 1120, 415),
        "storage": (790, 460, 1120, 585),
        "deploy": (1230, 230, 1450, 380),
    }
    draw_box(draw, boxes["user"], "Student / Admin", "Browser")
    draw_box(draw, boxes["ui"], "React interface", "Responsive UI", LIGHT_TEAL)
    draw_box(draw, boxes["state"], "Application shell", "Navigation + shared state", LIGHT_GRAY)
    draw_box(draw, boxes["screens"], "Feature screens", "Student and admin")
    draw_box(draw, boxes["models"], "TypeScript models", "Resource, Booking, Waitlist")
    draw_box(draw, boxes["storage"], "Local persistence", "Browser localStorage", LIGHT_GRAY)
    draw_box(draw, boxes["deploy"], "Vercel", "Static deployment", LIGHT_TEAL)
    draw_arrow(draw, (280, 305), (380, 180))
    draw_arrow(draw, (280, 305), (380, 435))
    draw_arrow(draw, (675, 180), (790, 160))
    draw_arrow(draw, (675, 435), (790, 345))
    draw_arrow(draw, (675, 435), (790, 520))
    draw_arrow(draw, (1120, 160), (1230, 285))
    draw_arrow(draw, (1120, 345), (1230, 320))
    image.save(path, quality=95)


def create_flow_diagram(path: Path, title_text: str, steps: Sequence[tuple[str, str]], accent=TEAL):
    width = 1500
    height = 240 + len(steps) * 155
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title = load_font(31, True)
    draw.text((70, 45), title_text, font=title, fill=f"#{DARK}")
    x1, x2 = 235, 1265
    y = 145
    for idx, (step_title, detail) in enumerate(steps):
        fill = LIGHT_TEAL if idx % 2 == 0 else LIGHT_GRAY
        draw.rounded_rectangle((x1, y, x2, y + 105), radius=18, fill=f"#{fill}", outline=f"#{accent}", width=3)
        number_font = load_font(23, True)
        body_font = load_font(22)
        draw.ellipse((x1 + 24, y + 22, x1 + 85, y + 83), fill=f"#{accent}")
        number = str(idx + 1)
        bbox = draw.textbbox((0, 0), number, font=number_font)
        draw.text((x1 + 54 - (bbox[2] - bbox[0]) / 2, y + 52 - (bbox[3] - bbox[1]) / 2 - 2), number, font=number_font, fill="white")
        title_font = load_font(23, True)
        draw.text((x1 + 110, y + 22), step_title, font=title_font, fill=f"#{DARK}")
        detail_lines = wrap_text(draw, detail, body_font, x2 - x1 - 145)
        for line_idx, line in enumerate(detail_lines[:2]):
            draw.text((x1 + 110, y + 56 + line_idx * 25), line, font=body_font, fill=f"#{MUTED}")
        if idx < len(steps) - 1:
            draw_arrow(draw, ((x1 + x2) // 2, y + 105), ((x1 + x2) // 2, y + 145), color=accent, width=4)
        y += 155
    image.save(path, quality=95)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 16, TEAL, 16, 8),
        (2, 13, DEEP_TEAL, 12, 6),
        (3, 11.5, DARK, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("SLOTWISE MMU  |  PROJECT DOCUMENTATION")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6500, 2860], indent_dxa=0)
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(table.cell(0, 0).paragraphs[0].add_run("Shortcut Asia Internship Challenge 2026"), size=8.5, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    set_run_font(right.add_run("Page "), size=8.5, color=MUTED)
    add_field(right, "PAGE")


def cover_page(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(20)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=0)
    cell = table.cell(0, 0)
    shade_cell(cell, TEAL)
    set_cell_margins(cell, top=420, start=420, bottom=420, end=420)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    set_run_font(p.add_run("SHORTCUT ASIA INTERNSHIP CHALLENGE 2026"), size=10, color=WHITE, bold=True)
    title = cell.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("SlotWise MMU"), size=32, color=WHITE, bold=True)
    subtitle = cell.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(0)
    set_run_font(subtitle.add_run("Campus Resource Booking System"), size=17, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    lead = doc.add_paragraph()
    lead.paragraph_format.space_after = Pt(16)
    lead.paragraph_format.line_spacing = 1.18
    set_run_font(
        lead.add_run("A responsive product prototype for discovering, booking, and managing shared university resources."),
        size=15,
        color=DARK,
        bold=True,
    )
    add_key_value_table(doc, [
        ("Candidate", "Karim Hamzaoui"),
        ("Submission", "Shortcut Asia Internship Challenge 2026"),
        ("Platform", "Responsive web application"),
        ("Technology", "React, TypeScript, Vite, Tailwind CSS, Vercel"),
        ("Date", "13 June 2026"),
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("LIVE APPLICATION"), size=9, color=CORAL, bold=True)
    link = doc.add_paragraph()
    link.paragraph_format.space_after = Pt(6)
    add_hyperlink(link, "https://slotwise-mmu.vercel.app", "https://slotwise-mmu.vercel.app")
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    set_run_font(p2.add_run("SOURCE CODE"), size=9, color=CORAL, bold=True)
    repo = doc.add_paragraph()
    add_hyperlink(repo, "github.com/karimhamzaoui05/slotwise-mmu", "https://github.com/karimhamzaoui05/slotwise-mmu")
    doc.add_page_break()


def add_contents(doc: Document):
    add_heading(doc, "Document guide", 1)
    add_body(doc, "This report explains the product problem, implemented functionality, architecture, key workflows, engineering decisions, limitations, roadmap, and setup process.")
    items = [
        "Executive summary and product context",
        "Target users, goals, and implemented features",
        "Booking and waitlist workflows",
        "Technical architecture and data model",
        "Engineering decisions, edge cases, and challenges",
        "Limitations, production roadmap, and setup instructions",
        "Suggested reviewer walkthrough",
    ]
    for item in items:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Review note",
        "The application is a functional front-end prototype. Authentication and shared persistence are simulated so reviewers can access the complete workflow without external credentials or infrastructure.",
    )


def build_document():
    DOCS.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    architecture = ASSETS / "architecture.png"
    booking_flow = ASSETS / "booking_flow.png"
    waitlist_flow = ASSETS / "waitlist_flow.png"
    create_architecture_diagram(architecture)
    create_flow_diagram(booking_flow, "Conflict-aware booking flow", [
        ("Discover", "Search or filter resources by type, building, capacity, and availability."),
        ("Inspect", "Review location, capacity, amenities, operating hours, and usage rules."),
        ("Select", "Choose a day and an available time slot from the weekly calendar."),
        ("Validate", "Reject unavailable, maintenance, or overlapping selections with clear feedback."),
        ("Confirm", "Create a booking reference, persist the booking, and show next actions."),
    ])
    create_flow_diagram(waitlist_flow, "Waitlist recovery flow", [
        ("Join queue", "Create a waitlist entry when the preferred resource or slot is unavailable."),
        ("Track demand", "Show queue position, likelihood, expiration, and the requested time."),
        ("Release slot", "Offer a cancelled slot to the first eligible student and start a countdown."),
        ("Claim or decline", "Convert a successful claim into a confirmed booking or offer it onward."),
    ], accent=GREEN)

    doc = Document()
    doc.core_properties.title = "SlotWise MMU Project Documentation"
    doc.core_properties.subject = "Shortcut Asia Internship Challenge 2026"
    doc.core_properties.author = "Karim Hamzaoui"
    doc.core_properties.keywords = "SlotWise MMU, booking system, React, TypeScript, internship challenge"
    configure_styles(doc)
    configure_page(doc)
    add_numbering_definition(doc, "bullet", "-", 31, 31)
    add_numbering_definition(doc, "decimal", "%1.", 32, 32)

    cover_page(doc)
    add_contents(doc)

    add_heading(doc, "1. Executive summary", 1)
    add_body(doc, "SlotWise MMU is a responsive campus resource booking application created for the Shortcut Asia Internship Challenge 2026. It gives students one place to discover and reserve rooms, equipment, laboratories, and project spaces, while giving administrators an operational view of utilization, maintenance, waitlists, and no-shows.")
    add_callout(doc, "Core value", "Make campus availability visible, prevent invalid bookings, and recover cancelled capacity through a structured waitlist.")
    add_key_value_table(doc, [
        ("Problem", "Fragmented or manual resource booking causes uncertainty, clashes, and wasted availability."),
        ("Solution", "A unified interface for discovery, conflict-aware booking, check-in, cancellation, waitlists, and administration."),
        ("Primary users", "MMU students and campus resource administrators."),
        ("Current scope", "Functional front-end prototype with browser-based persistent demo data."),
        ("Deployment", "Public Vercel application connected to the GitHub main branch."),
    ])

    add_heading(doc, "2. Problem and target users", 1)
    add_heading(doc, "Problem statement", 2)
    add_body(doc, "Campus resources are often coordinated through separate forms, spreadsheets, messages, or manual enquiries. Students may not know what is available, cancelled slots may remain unused, and administrators may lack a clear view of demand and operational status.")
    for text in (
        "Students cannot quickly compare availability, capacity, location, and amenities.",
        "Manual processes make double-booking and overlapping reservations harder to prevent.",
        "Cancelled reservations can leave useful rooms or equipment idle.",
        "Waitlist position, check-in status, and upcoming reservations are difficult to track.",
        "Administrators need one view of occupancy, maintenance, waitlists, and no-shows.",
    ):
        add_bullet(doc, text)
    add_heading(doc, "Target users", 2)
    add_comparison_table(doc, ["User", "Primary need", "Supported actions"], [
        ("Students", "Find and reserve the right resource with minimal friction.", "Search, filter, inspect, book, check in, edit, cancel, join and claim waitlists."),
        ("Administrators", "Monitor operations and keep resources accurate and available.", "Review metrics, utilization and activity; add, edit, block, maintain, or deactivate resources."),
    ], [1700, 3000, 4660])

    add_heading(doc, "3. Product goals and feature scope", 1)
    goals = [
        "Make availability immediately understandable.",
        "Prevent invalid, unavailable, maintenance, and overlapping bookings.",
        "Reduce wasted capacity through waitlist recovery.",
        "Support efficient repeated use on desktop and mobile.",
        "Demonstrate complete product states, not only a happy-path form.",
    ]
    for goal in goals:
        add_numbered(doc, goal)
    add_heading(doc, "Implemented feature set", 2)
    add_comparison_table(doc, ["Area", "Implemented details"], [
        ("Demo access", "Student and Admin shortcuts, university email validation UI, role-based navigation."),
        ("Dashboard", "Search, quick availability controls, available-now resources, bookings, occupancy, and waitlist alerts."),
        ("Discovery", "Search, filters, sorting, grid/list views, status, location, capacity, amenities, and next availability."),
        ("Booking", "Weekly calendar, time-slot states, conflict feedback, live summary, booking reference, and confirmation."),
        ("My Bookings", "Upcoming, past, and cancelled tabs with check-in, edit, cancellation, and direction actions."),
        ("Waitlist", "Queue position, likelihood, expiration, countdown, claim, decline, and leave actions."),
        ("Administration", "Operational metrics, charts, schedule, activity, resource table, edit drawer, and maintenance blocks."),
        ("Responsive UX", "Desktop sidebar, mobile bottom navigation, accessible labels, focus states, and touch targets."),
    ], [2100, 7260])

    doc.add_page_break()
    add_heading(doc, "4. Key product workflows", 1)
    add_heading(doc, "Booking workflow", 2)
    add_body(doc, "The booking experience goes deeper than basic create-and-save functionality. It communicates resource rules, prevents invalid selections, demonstrates conflict handling, and keeps the booking consistent across confirmation and management screens.")
    doc.add_picture(str(booking_flow), width=Inches(6.35))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    set_run_font(caption.add_run("Figure 1. Conflict-aware student booking flow"), size=9, color=MUTED, italic=True)
    add_callout(doc, "Overlap rule", "A production conflict check uses: existingStart < requestedEnd AND requestedStart < existingEnd. This permits back-to-back bookings while rejecting intersecting ranges.", fill="FFF7ED", accent=GOLD)
    add_heading(doc, "Waitlist workflow", 2)
    add_body(doc, "The waitlist converts cancellations into usable capacity. Students see their position and a released slot countdown; a successful claim removes the queue entry and creates a confirmed booking.")
    doc.add_picture(str(waitlist_flow), width=Inches(6.35))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    set_run_font(caption.add_run("Figure 2. Waitlist release and claim flow"), size=9, color=MUTED, italic=True)

    doc.add_page_break()
    add_heading(doc, "5. Technical architecture", 1)
    add_body(doc, "SlotWise MMU is a client-side React application compiled by Vite and deployed as static assets on Vercel. Shared booking and waitlist state is managed at the application shell and persisted in the browser for a reliable reviewer demo.")
    doc.add_picture(str(architecture), width=Inches(6.35))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    set_run_font(caption.add_run("Figure 3. Current prototype architecture"), size=9, color=MUTED, italic=True)
    add_heading(doc, "Technology stack", 2)
    add_comparison_table(doc, ["Technology", "Purpose"], [
        ("React 18", "Component-based interface and stateful product interactions."),
        ("TypeScript", "Typed contracts for screens, resources, bookings, statuses, and waitlist entries."),
        ("Vite", "Fast development server and optimized production build."),
        ("Tailwind CSS", "Responsive utility styling supported by shared design tokens."),
        ("Lucide React", "Consistent interface icons and accessible visual cues."),
        ("Recharts", "Administrative utilization and booking trend visualizations."),
        ("Sonner", "Clear success and status notifications."),
        ("Vercel", "Public deployment and automatic redeployment from GitHub."),
    ], [2350, 7010])
    add_heading(doc, "Application structure", 2)
    add_comparison_table(doc, ["Location", "Responsibility"], [
        ("src/app/App.tsx", "Navigation, user role, shared booking/waitlist state, events, and local persistence."),
        ("src/app/components/screens", "Dashboard, Explore, Resource Details, Booking, Waitlist, Admin, and Profile screens."),
        ("src/app/components", "Reusable navigation, resource cards, top bars, and status badges."),
        ("src/app/types.ts", "Resource, Booking, WaitlistEntry, User, status, and screen contracts."),
        ("src/app/data/mockData.ts", "Realistic MMU-style resources, bookings, waitlists, and admin metrics."),
        ("src/styles", "Global presentation, design tokens, responsive behavior, and accessibility rules."),
    ], [2850, 6510])

    add_heading(doc, "6. Technical decisions", 1)
    decisions = [
        ("React with TypeScript", "The interface contains reusable, stateful workflows. TypeScript improves clarity and catches invalid values or incomplete transitions."),
        ("Vite for the prototype", "The challenge prioritizes a dependable product within a short period. Vite keeps the build simple and fast without unnecessary server complexity."),
        ("Central shared state", "Bookings and waitlists live at the application shell so actions on one screen immediately update confirmation, My Bookings, and Waitlist views."),
        ("Browser persistence", "localStorage keeps the demo state after refresh and avoids database setup for reviewers. It is deliberately separated from the proposed production architecture."),
        ("Responsive web delivery", "Students can book from phones while administrators can work from desktops, using one deployable codebase."),
        ("Reusable components", "Navigation, resource cards, status badges, and design tokens improve consistency and reduce duplicated behavior."),
    ]
    for title, detail in decisions:
        add_body(doc, f"{title}: {detail}", bold_prefix=f"{title}:")

    add_heading(doc, "7. Edge cases and quality considerations", 1)
    for text in (
        "Unavailable and maintenance slots cannot be selected.",
        "A demonstrated overlap state provides immediate inline conflict feedback.",
        "Cancellation requires confirmation to reduce accidental destructive actions.",
        "Upcoming, past, cancelled, and waitlist screens include intentional empty states.",
        "Duplicate waitlist entries for the same resource are detected.",
        "Claimed waitlist slots become confirmed bookings in shared state.",
        "Malformed browser storage falls back to known mock data.",
        "Mobile layouts were verified at 390 px without horizontal overflow.",
        "Status is communicated through text and icons rather than color alone.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "8. Challenges and solutions", 1)
    add_comparison_table(doc, ["Challenge", "Resolution"], [
        ("Turning a design into maintainable code", "Separated application state, typed models, reusable components, and screen-level responsibilities instead of isolated mock screens."),
        ("Keeping state consistent", "Lifted bookings and waitlists into the application shell and used typed event handlers for create, update, leave, and claim operations."),
        ("Safely parsing time keys", "Removed only the day prefix from values such as Fri:10:00, preserving the complete time rather than splitting on every colon."),
        ("Balancing scope and depth", "Focused on two connected workflows: conflict-aware booking and waitlist recovery, supported by a useful admin experience."),
        ("Making review easy", "Used realistic sample data, demo role shortcuts, public deployment, and browser persistence with no required environment variables."),
    ], [3100, 6260])

    add_heading(doc, "9. Automated testing", 1)
    add_body(doc, "Core booking rules are extracted into pure TypeScript functions and verified with Vitest independently from the React interface.")
    for text in (
        "Partial and contained booking overlaps.",
        "Valid back-to-back bookings.",
        "Conflict filtering by resource and date.",
        "Waitlist claim conversion and queue removal.",
        "Rejection of unreleased waitlist slots.",
        "Immutable booking cancellation.",
    ):
        add_bullet(doc, text)
    add_callout(doc, "Test command", "Run npm test to execute the complete unit suite.")

    add_heading(doc, "10. Current limitations", 1)
    add_callout(doc, "Prototype boundary", "The application is intentionally optimized for evaluation and demonstration. The limitations below are explicit, not hidden production claims.", fill="FFF1F2", accent=RED)
    for text in (
        "Authentication is simulated and does not verify a real MMU identity.",
        "Data is stored per browser and is not synchronized across users or devices.",
        "Resource and administration data starts from realistic mock data.",
        "Email, push notifications, and calendar integration are not connected to external services.",
        "The final conflict check is demonstrated client-side rather than enforced by a transactional database.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "11. Production roadmap", 1)
    roadmap = [
        "Add Supabase or PostgreSQL for shared persistent data.",
        "Implement MMU email authentication and role-based authorization.",
        "Enforce booking overlap rules with database constraints or transactions.",
        "Expand the unit suite with React interaction and end-to-end browser tests.",
        "Add server-generated reminders and waitlist release notifications.",
        "Support QR-code check-in, calendar export, audit logs, and utilization reports.",
    ]
    for item in roadmap:
        add_numbered(doc, item)

    doc.add_page_break()
    add_heading(doc, "12. Setup, build, and deployment", 1)
    add_heading(doc, "Requirements", 2)
    for item in ("Node.js LTS", "npm", "A modern browser such as Chrome, Edge, Firefox, or Safari"):
        add_bullet(doc, item)
    add_heading(doc, "Run locally", 2)
    code_lines = [
        "git clone https://github.com/karimhamzaoui05/slotwise-mmu.git",
        "cd slotwise-mmu",
        "npm install",
        "npm run dev",
    ]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade_cell(cell, DARK)
    set_cell_margins(cell, top=180, start=200, bottom=180, end=200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code_lines):
        if idx:
            p.add_run("\n")
        set_run_font(p.add_run(line), size=10, color=WHITE, name="Consolas")
    add_body(doc, "Open the URL printed by Vite, normally http://localhost:5173. Select Student Demo or Admin Demo on the sign-in screen.", after=8)
    add_heading(doc, "Quality commands", 2)
    add_key_value_table(doc, [
        ("Automated tests", "npm test"),
        ("Type-check", "npm run typecheck"),
        ("Production build", "npm run build"),
        ("Preview build", "npm run preview"),
    ])
    add_heading(doc, "Deployment", 2)
    add_body(doc, "The GitHub repository is connected to Vercel. A successful push to the main branch triggers a new production deployment. The current prototype requires no environment variables.")
    p = doc.add_paragraph()
    set_run_font(p.add_run("Live application: "), bold=True)
    add_hyperlink(p, "https://slotwise-mmu.vercel.app", "https://slotwise-mmu.vercel.app")
    p2 = doc.add_paragraph()
    set_run_font(p2.add_run("GitHub repository: "), bold=True)
    add_hyperlink(p2, "github.com/karimhamzaoui05/slotwise-mmu", "https://github.com/karimhamzaoui05/slotwise-mmu")

    add_heading(doc, "13. Suggested reviewer walkthrough", 1)
    walkthrough = [
        "Open the live application and choose Student Demo.",
        "Review the dashboard and available resources.",
        "Open Explore Resources and test search, filters, and view modes.",
        "Open a resource, test the conflict state, and reserve a valid slot.",
        "Confirm that the new reservation appears in My Bookings.",
        "Open Waitlist and claim the released Collaboration Room slot.",
        "Sign out and choose Admin Demo.",
        "Review metrics, charts, Resource Management, and maintenance actions.",
    ]
    for item in walkthrough:
        add_numbered(doc, item)

    add_heading(doc, "14. Conclusion", 1)
    add_body(doc, "SlotWise MMU demonstrates a complete product workflow rather than a collection of static screens. It addresses a realistic campus problem, includes meaningful booking and waitlist logic, handles important edge cases, and supports both student and administrative users. The architecture keeps the prototype easy to run and review while providing a clear path toward a secure, multi-user production system.")
    add_callout(doc, "Submission position", "A working public application, transparent technical documentation, and a focused product story built around depth, usability, and explainable engineering decisions.")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
